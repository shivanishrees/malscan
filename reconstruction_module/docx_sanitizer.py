import zipfile
from docx import Document

def sanitize_docx(source, dest):
    # macro removal happens because python-docx doesn’t include VBA
    doc = Document(source)
    clean = Document()

    for para in doc.paragraphs:
        clean.add_paragraph(para.text)

    clean.save(dest)
